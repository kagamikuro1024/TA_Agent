import os
import shutil
from docx import Document
from pypdf import PdfWriter

# Cấu hình đường dẫn
ROOT_DIR = os.getcwd()
MOCK_DIR = os.path.join(ROOT_DIR, "tests", "stress_data")
SOURCE_DATA = os.path.join(ROOT_DIR, "data")
HUGE_PDF_SOURCE = os.path.join(SOURCE_DATA, "en_CCNAS_v11_Ch01-Modern Network Security Threats.pdf")

def setup_dirs():
    if not os.path.exists(MOCK_DIR):
        os.makedirs(MOCK_DIR)
        print(f"Created directory: {MOCK_DIR}")

def create_corrupt_pdf():
    path = os.path.join(MOCK_DIR, "corrupt_header.pdf")
    with open(path, "wb") as f:
        f.write(b"This is a fake PDF content. Not a real binary header.\n" * 100)
    print(f"Created corrupt PDF: {path}")

def create_complex_docx():
    path = os.path.join(MOCK_DIR, "complex_structure.docx")
    doc = Document()
    doc.add_heading('Hợp đồng Điện tử AIK-024', 0)
    
    doc.add_heading('1. Tiêu đề H1', level=1)
    doc.add_paragraph('Đây là nội dung của chương 1.')
    
    doc.add_heading('1.1. Tiêu đề H2', level=2)
    p = doc.add_paragraph('Danh sách các yêu cầu:')
    doc.add_paragraph('Yêu cầu 1', style='List Bullet')
    doc.add_paragraph('Yêu cầu 2', style='List Bullet')
    
    doc.add_heading('1.1.1. Tiêu đề H3', level=3)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"Ô ({r},{c})"
            
    doc.add_heading('1.1.1.1. Tiêu đề H4', level=4)
    doc.add_paragraph('Ký tự đặc biệt: → ≥ α β γ')
    
    doc.save(path)
    print(f"Created complex DOCX: {path}")

def create_empty_docx():
    path = os.path.join(MOCK_DIR, "empty_xml.docx")
    doc = Document()
    doc.save(path)
    print(f"Created empty DOCX: {path}")

def create_password_pdf():
    path = os.path.join(MOCK_DIR, "password_locked.pdf")
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.encrypt("mật-khẩu-bí-mật")
    with open(path, "wb") as f:
        writer.write(f)
    print(f"Created password locked PDF: {path}")

def copy_huge_pdf():
    dest = os.path.join(MOCK_DIR, "huge_slide.pdf")
    if os.path.exists(HUGE_PDF_SOURCE):
        shutil.copy(HUGE_PDF_SOURCE, dest)
        print(f"Copied huge PDF from {HUGE_PDF_SOURCE} to {dest}")
    else:
        # Fallback if file not found
        print(f"Warning: Source huge PDF not found at {HUGE_PDF_SOURCE}. Creating a fallback large file.")
        with open(dest, "wb") as f:
            f.write(b"%PDF-1.5\n" + os.urandom(30 * 1024 * 1024))

if __name__ == "__main__":
    setup_dirs()
    create_corrupt_pdf()
    create_complex_docx()
    create_empty_docx()
    create_password_pdf()
    copy_huge_pdf()
    print("Mock data generation complete.")
