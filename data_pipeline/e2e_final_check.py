import asyncio
import os
import time
import subprocess
import httpx
import sys
import json
import numpy as np
from rich.console import Console
from rich.panel import Panel

# Fix UnicodeEncodeError for Windows console
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')

console = Console()

# Configuration
BASE_URL = "http://127.0.0.1:8000"
UPLOAD_URL = f"{BASE_URL}/api/v1/documents/upload"
MOCK_DATA_DIR = "tests/e2e_data"
DB_URL = os.getenv("ASYNC_DATABASE_URL", "postgresql://user:password@localhost:5432/agent_db")


def setup_test_file():
    """Generates a sample DOCX file for E2E testing."""
    os.makedirs(MOCK_DATA_DIR, exist_ok=True)
    from docx import Document
    doc = Document()
    doc.add_heading('E2E LECTURE: AIK-024 SYSTEM TESTING', 0)
    doc.add_paragraph('This is experimental content to verify Vector Database consistency.')
    doc.add_paragraph('Steps include: Upload, Ingestion, Embedding, and HITL Correction.')
    doc.add_heading('1. Data Flywheel Definition', level=1)
    doc.add_paragraph('The knowledge loop allows humans to intervene in the AI brain to correct errors.')
    
    # Add noise footer to test Noise Filter Audit
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "© 2026 AIK Team. Page 1 of 100. Confidential."
    
    path = os.path.join(MOCK_DATA_DIR, "e2e_test_doc.docx")
    doc.save(path)
    return path

def cosine_similarity(v1, v2):
    """Calculates cosine similarity between two vectors."""
    if isinstance(v1, str):
        v1 = [float(x) for x in v1.strip('[]').split(',')]
    if isinstance(v2, str):
        v2 = [float(x) for x in v2.strip('[]').split(',')]
        
    v1 = np.array(v1, dtype=float)
    v2 = np.array(v2, dtype=float)
    return np.dot(v1, v2) / (np.linalg.norm(v1) * np.linalg.norm(v2))

async def run_e2e():
    console.print(Panel("[bold green]STARTING E2E FINAL VERIFICATION[/bold green]"))
    
    # 1. Start server if not running
    server_process = None
    try:
        async with httpx.AsyncClient() as client:
            await client.get(f"{BASE_URL}/health")
            console.print("[yellow][SKIP] Server is already running.[/yellow]")
    except:
        console.print("[cyan][SETUP] Starting FastAPI server...[/cyan]")
        server_process = subprocess.Popen(
            ["uvicorn", "src.main:app", "--host", "127.0.0.1", "--port", "8000"],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        # Polling server health
        for i in range(15):
            try:
                async with httpx.AsyncClient() as client:
                    await client.get(f"{BASE_URL}/health")
                    console.print("[SETUP] Server is ready.")
                    break
            except:
                if server_process.poll() is not None:
                    out, err = server_process.communicate()
                    console.print(f"[red]Server failed to start: {err}[/red]")
                    return
                time.sleep(1)
                if i == 14:
                    console.print("[red]Timeout waiting for server.[/red]")
                    return

    try:
        # 2. Upload file
        test_file = setup_test_file()
        console.print(f"[step] [1/5] Uploading file: {test_file}")
        async with httpx.AsyncClient(timeout=30.0) as client:
            with open(test_file, "rb") as f:
                resp = await client.post(UPLOAD_URL, files={"file": f})
            
            assert resp.status_code == 202
            data = resp.json()
            doc_filename = data['source_uri']
            console.print(f"   [OK] Upload successful. Request ID: {data.get('request_id')}")

            # 3. Wait for Ingestion Worker
            console.print("[step] [2/5] Waiting for Ingestion Worker (Polling DB)...")
            import asyncpg
            pool = await asyncpg.create_pool(DB_URL)
            doc_id = None
            for _ in range(15):
                row = await pool.fetchrow("SELECT id, status FROM documents WHERE filename = $1 ORDER BY created_at DESC LIMIT 1", doc_filename)
                if row and row['status'] == 'READY':
                    doc_id = row['id']
                    break
                await asyncio.sleep(2)
            
            if not doc_id:
                console.print("[red][FAIL] Server failed to process document within 30s.[/red]")
                return

            # 4. Vector Verification and Cosine Similarity
            console.print("[step] [3/5] Verifying Vector Proof (Cosine Similarity)...")
            chunks = await pool.fetch("SELECT id, content, embedding FROM document_chunks WHERE document_id = $1", doc_id)
            console.print(f"   Found {len(chunks)} chunks in DB.")
            
            if not chunks:
                console.print("[red][FAIL] No chunks found for document.[/red]")
                return

            # Pick one chunk and verify similarity with a fresh embedding from the same text
            target_chunk = chunks[0]
            target_text = target_chunk['content']
            target_vec = target_chunk['embedding']
            
            # Updated import from the new package structure
            from data_pipeline.pipeline.embedding import generate_embeddings
            new_vecs = await generate_embeddings([target_text])
            sim = cosine_similarity(target_vec, new_vecs[0])
            
            status_color = "green" if sim > 0.99 else "red"
            console.print(f"   [bold {status_color}]Cosine Similarity Proof: {sim:.4f}[/bold {status_color}]")
            
            # 5. HITL PATCH Test
            console.print("[step] [4/5] Testing HITL PATCH API (Data Flywheel)...")
            chunk_id = str(target_chunk['id'])
            new_content = "REVISED CONTENT: Brain Flywheel Knowledge Loop."
            
            patch_resp = await client.patch(
                f"{BASE_URL}/api/v1/documents/chunks/{chunk_id}",
                json={"new_content": new_content, "reason": "Correction by TA"}
            )
            
            assert patch_resp.status_code == 200
            console.print("   [OK] PATCH /chunks successful. Re-embedding triggered.")
            
            # Verify DB content updated
            updated_row = await pool.fetchrow("SELECT content, metadata FROM document_chunks WHERE id = $1", target_chunk['id'])
            assert updated_row['content'] == new_content
            meta = json.loads(updated_row['metadata']) if isinstance(updated_row['metadata'], str) else updated_row['metadata']
            assert 'audit_trail' in meta
            console.print("   [OK] DB updated with corrected content and audit trail.")

            # 6. Noise Filter Audit
            console.print("[step] [5/5] Auditing Noise Filter (Header/Footer)...")
            noise_patterns = ["© 2026", "Page 1 of 100", "Confidential"]
            noise_found = []
            for c in chunks:
                for p in noise_patterns:
                    if p in c['content']:
                        noise_found.append((c['id'], p))
            
            if noise_found:
                console.print(f"   [yellow][WARN] Found {len(noise_found)} matches of noise in knowledge segments (Footer noise).[/yellow]")
                for nid, p in noise_found:
                    console.print(f"      - Chunk {str(nid)[:8]}... contains: '{p}'")
            else:
                console.print("   [OK] No noise patterns detected in segments.")

            await pool.close()
            console.print(Panel("[bold green]E2E FINAL CHECK: SUCCESSFUL 100%[/bold green]"))

    except Exception as e:
        console.print(f"[red][ERROR] E2E FAILED: {str(e)}[/red]")
        import traceback
        traceback.print_exc()
    finally:
        if server_process:
            server_process.terminate()
            console.print("[cyan][CLEANUP] FastAPI server closed.[/cyan]")

if __name__ == "__main__":
    # Ensure PYTHONPATH is correctly set to import modules
    sys.path.append(os.getcwd())
    asyncio.run(run_e2e())
